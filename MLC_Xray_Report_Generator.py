"""
MLC X-Ray Report Generator (Streamlit app)

Usage:
1. Install dependencies:
   pip install streamlit fpdf python-docx Pillow
2. Run:
   streamlit run MLC_Xray_Report_Generator.py

This creates a simple web UI to enter patient/hospital details, X-ray findings,
and generates a downloadable PDF report. You can also export a DOCX.
"""

from io import BytesIO
from datetime import datetime
import streamlit as st
from fpdf import FPDF
from PIL import Image
from docx import Document

st.set_page_config(page_title="MLC X-ray Report Generator", layout="centered")

st.title("MLC X-ray Report Generator")
st.markdown("Simple tool to create and download MLC X-ray reports as PDF or DOCX.")

# --- Patient / Case details ---
with st.form("details_form"):
    col1, col2 = st.columns(2)
    with col1:
        patient_name = st.text_input("Patient Name", "")
        age = st.text_input("Age", "")
        sex = st.selectbox("Sex", ["", "Male", "Female", "Other"])
        hospital_no = st.text_input("Hospital / OPD No.", "")
    with col2:
        referring_physician = st.text_input("Referring Physician", "")
        date_of_exam = st.date_input("Date of Exam", datetime.today())
        xray_type = st.selectbox("X-ray Type", ["Chest PA", "Chest AP", "Skull", "KUB", "Spine", "Other"]) 
        if xray_type == "Other":
            xray_type = st.text_input("Specify X-ray Type", "")

    clinical_history = st.text_area("Clinical History / Complaint", "")
    findings = st.text_area("Findings (describe radiological findings)", "")
    impression = st.text_area("Impression / Conclusion", "")
    doctor_name = st.text_input("Reporting Doctor / Radiologist", "")
    signature_img = st.file_uploader("Upload signature image (optional)", type=["png","jpg","jpeg"]) 

    generate = st.form_submit_button("Generate Report")

# Helper: create a nicely formatted report text
def build_report_text():
    lines = []
    lines.append("MLC X-RAY REPORT")
    lines.append("")
    lines.append(f"Patient Name: {patient_name}")
    lines.append(f"Age / Sex: {age} / {sex}")
    lines.append(f"Hospital / OPD No.: {hospital_no}")
    lines.append(f"Referring Physician: {referring_physician}")
    lines.append(f"Date of Exam: {date_of_exam.strftime('%d-%m-%Y')}")
    lines.append(f"Examination: {xray_type}")
    lines.append("")
    if clinical_history:
        lines.append("Clinical History:")
        lines.extend([clinical_history])
        lines.append("")
    lines.append("Findings:")
    lines.extend([findings if findings else "-"])
    lines.append("")
    lines.append("Impression:")
    lines.extend([impression if impression else "-"])
    lines.append("")
    lines.append(f"Reporting Doctor: {doctor_name}")
    lines.append("")
    lines.append("Note: This report is generated electronically and is valid without a wet signature unless otherwise required.")
    return "\n".join(lines)

# PDF generation using FPDF
class PDF(FPDF):
    def header(self):
        pass

    def footer(self):
        self.set_y(-20)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Generated on {datetime.now().strftime("%d-%m-%Y %H:%M")}', 0, 0, 'C')


def create_pdf(report_text, signature_bytes=None):
    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, 'MLC X-RAY REPORT', ln=True, align='C')
    pdf.ln(4)
    pdf.set_font('Arial', '', 11)

    for line in report_text.split('\n'):
        pdf.multi_cell(0, 7, line)
    pdf.ln(6)

    if signature_bytes:
        try:
            img = Image.open(BytesIO(signature_bytes))
            with BytesIO() as img_buf:
                img.save(img_buf, format='PNG')
                img_bytes = img_buf.getvalue()
            pdf.image(BytesIO(img_bytes), x=150, y=pdf.get_y(), w=40)
        except Exception:
            pass

    out = BytesIO()
    pdf.output(out)
    out.seek(0)
    return out

# DOCX generation
def create_docx(report_text, signature_bytes=None):
    doc = Document()
    doc.add_heading('MLC X-RAY REPORT', level=1)
    for para in report_text.split('\n'):
        doc.add_paragraph(para)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

if generate:
    report_text = build_report_text()

    sig_bytes = None
    if signature_img:
        sig_bytes = signature_img.read()

    pdf_file = create_pdf(report_text, signature_bytes=sig_bytes)

    st.success("Report generated — download below")

    with st.expander("Preview report text"):
        st.text(report_text)

    st.download_button("Download PDF", data=pdf_file, file_name=f"MLC_Xray_Report_{patient_name or 'patient'}.pdf", mime='application/pdf')

    try:
        docx_file = create_docx(report_text, signature_bytes=sig_bytes)
        st.download_button("Download DOCX", data=docx_file, file_name=f"MLC_Xray_Report_{patient_name or 'patient'}.docx", mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception:
        pass

    st.markdown("**Copy report text** — select below and copy (Ctrl+C / Cmd+C):")
    st.code(report_text)

else:
    st.info("Fill patient details and click 'Generate Report' to create PDF/DOCX.")

st.markdown("---")
st.markdown("**Tips:** Customize the report wording in the script. For hospital branding, you can add a logo in the PDF header section of the PDF class.")
